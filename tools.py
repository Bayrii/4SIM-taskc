from docx import Document
import re


def write_to_docx(text):
    """Writing text to .docx"""
    print("Creating docx file...\n")

    doc = Document()

    # Başlıqları və mətni ayır
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue

        # Başlıq (tam sətir böyük hərflərlə və ulduzsuz)
        if re.fullmatch(r"\*{0,2}[A-ZƏÖÜĞÇŞİ\s]+", line):
            clean_title = line.replace("*", "").strip()
            doc.add_heading(clean_title, level=1)

        # Bullet point
        elif line.startswith("* "):
            bullet_text = line[2:].strip()
            # Qalın sözləri formatlamaq üçün
            paragraph = doc.add_paragraph(style='List Bullet')
            while "**" in bullet_text:
                before, bold, after = re.split(r"\*\*(.+?)\*\*", bullet_text, 1)
                paragraph.add_run(before)
                paragraph.add_run(bold).bold = True
                bullet_text = after
            paragraph.add_run(bullet_text)

        # Adi mətn
        else:
            paragraph = doc.add_paragraph()
            while "**" in line:
                before, bold, after = re.split(r"\*\*(.+?)\*\*", line, 1)
                paragraph.add_run(before)
                paragraph.add_run(bold).bold = True
                line = after
            paragraph.add_run(line)

    doc.save(r"output.docx")

def write_to_txt(text):
    """Writing parsed text to .txt"""
    print("Creating txt file...\n")

    formated_headers_text = re.sub(r'\*\*([A-ZƏÖÜĞÇŞİ\s]+)\*\*', r'\1', text)

    filepath = r"output.txt"

    with open(filepath, "w", encoding="utf-8") as f:
        f.write(formated_headers_text)
        
    print("File saved as output.txt")
    
        # for line in text.split("\n"):
        #     line = line.strip()
        #     if not line:
        #         continue

        
        #     if re.fullmatch(r"\*{0,2}[A-ZƏÖÜĞÇŞİ\s]+", line):
        #         clean_title = line.replace("*", "").strip()
        #         f.write(clean_title + "\n\n")
        #     else:
        #         f.write(line)




def write_to_html(text):

    html_output = convert_to_html(text)

    # Save to HTML file
    with open('report.html', 'w', encoding='utf-8') as f:
        f.write(html_output)


def convert_to_html(text):
    html_lines = []

    # HTML header + CSS styling
    html_lines.append("""
<!DOCTYPE html>
<html lang="az">
<head>
<meta charset="UTF-8">
<title>Rəy və Təklif Asistantı</title>
<style>
body {
    font-family: "Segoe UI", Tahoma, Geneva, Verdana, sans-serif;
    background-color: #f4f6f8;
    padding: 30px;
    color: #333;
    line-height: 1.65;
}

.section {
    background-color: #ffffff;
    border-left: 8px solid #4CAF50;
    padding: 20px 25px;
    margin-bottom: 25px;
    border-radius: 10px;
    box-shadow: 0 4px 10px rgba(0,0,0,0.07);
    transition: transform 0.2s;
}

.section:hover {
    transform: translateY(-2px);
}

.section h2 {
    margin-top: 0;
    margin-bottom: 15px;
    color: #4CAF50;
    font-size: 1.5em;
    letter-spacing: 0.5px;
    border-bottom: 2px solid #e0e0e0;
    padding-bottom: 5px;
}

ul {
    padding-left: 20px;
}

li {
    margin-bottom: 12px;
    position: relative;
    padding-left: 10px;
}

li::before {
    content: "•";
    color: #4CAF50;
    position: absolute;
    left: 0;
}

p {
    margin: 12px 0;
    color: #444;
}

strong {
    color: #222;
    font-weight: 600;
}
</style>
</head>
<body>
""")

    section_open = False

    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue

        # Convert inline **bold** anywhere in the line
        line = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', line)

        # Headers: full line is uppercase letters (A-ZƏÖÜĞÇŞİ and spaces)
        if re.fullmatch(r'[A-ZƏÖÜĞÇŞİ0-9\s]+', line):
            if section_open:
                html_lines.append("</ul></div>")
                section_open = False
            html_lines.append(f'<div class="section"><h2>{line}</h2><ul>')
            section_open = True

        # Bullets or numbered items
        elif re.match(r'^(\d+\.|\*)\s+', line):
            content = re.sub(r'^(\d+\.|\*)\s+', '', line)
            html_lines.append(f'<li>{content}</li>')

        # Regular paragraph
        else:
            html_lines.append(f'<p>{line}</p>')

    if section_open:
        html_lines.append("</ul></div>")

    html_lines.append("</body></html>")
    return "\n".join(html_lines)
