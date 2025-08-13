import google.generativeai as genai
from google.generativeai.types import GenerationConfig
from docx import Document
import re
from prompt import build_prompt
from test_file import extract_text



def get_response(prompt_, mode_name='gemini-1.5-flash'):
    """Getting response"""
    prompt = prompt_

    model = genai.GenerativeModel(
        model_name=mode_name,
        system_instruction="Siz 'Rəy və Təklif Asistantı' — Azərbaycan Respublikası dövlət idarəetmə orqanlarında və hüquq sistemində"
                           "25 ildən artıq zəngin və çoxşaxəli təcrübəyə malik, strateji düşüncə qabiliyyəti yüksək, müasir hüquqi"
                           "normativləri dərindən mənimsəmiş, nüfuzlu və tanınmış ali səviyyəli dövlət analitiki və hüquq eksperti kimi"
                           "fəaliyyət göstərirsiniz."
    )

    generation_config = GenerationConfig(
        temperature=0.2
    )


    try:

        response = model.generate_content(
            contents=[
                {"role": "user", "parts": [{"text": prompt}]}
            ],
            generation_config=generation_config
        )

        if response.candidates and response.candidates[0].content and response.candidates[0].content.parts:
            return response.candidates[0].content.parts[0].text
        else:
            print("Error: Model did not return expected content structure.")
            return "Content generation failed."


    except Exception as e:
        print('An error occurred during content generation:' + e.__str__())
        return "Error: " + e.__str__()


def parse_response(response_text):
    """Parsing response"""
    text = """**ÜMUMİ RƏY**

                "Rəqəmsal İnnovasiya və Kibertəhlükəsizlik Şöbəsi"nin yaradılması təklifi, Azərbaycanın rəqəmsal transformasiya strategiyasına uyğun olaraq, əhəmiyyətli və zəruri bir addımdır. Lakin, təklifin maliyyə və təşkilati aspektləri daha dəqiq araşdırılmalı, həmçinin mövcud qanunvericiliyə tam uyğunluğunun təmin edilməsi vacibdir.
                
                **REGULYATİV VƏ STRATEJİ UYĞUNLUQ**
                
                Təklif "Rəqəmsal Azərbaycan" strategiyasının məqsədlərinə uyğundur və kibertəhlükəsizlik sahəsində mövcud qanunvericiliyə (məsələn, "İnformasiya təhlükəsizliyi haqqında" Qanun) əsaslanır.  Lakin, nazirliyin strukturuna yeni şöbənin əlavə edilməsi ilə bağlı "Dövlət qulluğu haqqında" Qanuna uyğunluğun təmin edilməsi və əməkdaşların seçimi üçün müvafiq prosedurların tətbiqi vacibdir.  Oxşar təcrübələr digər dövlət qurumlarının təcrübələrini araşdırmaqla müqayisə edilməlidir.  Məsələn, Rabitə Nazirliyinin kibertəhlükəsizlik strukturu və fəaliyyəti ilə müqayisəli təhlil aparılmalıdır.  Ziddiyyətli məqamlar aşkar edilərsə, uyğunlaşdırma tədbirləri görülməlidir.
            """

    parts = {
        "ÜMUMİ RƏY" : "",
        "REGULYATİV VƏ STRATEJİ UYĞUNLUQ": "",
        "GÜCLÜ TƏRƏFLƏR": "",
        "ZƏİF VƏ PROBLEMLİ SAHƏLƏR": "",
        "TƏKLİFLƏR": "",
        "HÜQUQİ QEYDLƏR": "",
        "MALİYYƏ ƏSASLARI": "",
        "YEKUN QİYMƏTLƏNDİRMƏ": ""
    }

    header = ""

    for ind, char in enumerate(text):

        pass
        # if header ==
        # pass

    pass

def write_to_docx(text):
    """Writing text to files"""

    doc = Document()

    # Başlıqları və mətni ayır
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue

        # Başlıq (tam sətir böyük hərflərlə və ulduzsuz)
        if re.fullmatch(r"\*{0,2}[A-ZƏÖÜĞÇŞİ\s]+", line):
            clean_title = line.replace("*", "").strip()
            doc.add_heading(clean_title, level=1)

        # Bullet point
        elif line.startswith("* "):
            bullet_text = line[2:].strip()
            # Qalın sözləri formatlamaq üçün
            paragraph = doc.add_paragraph(style='List Bullet')
            while "**" in bullet_text:
                before, bold, after = re.split(r"\*\*(.+?)\*\*", bullet_text, 1)
                paragraph.add_run(before)
                paragraph.add_run(bold).bold = True
                bullet_text = after
            paragraph.add_run(bullet_text)

        # Adi mətn
        else:
            paragraph = doc.add_paragraph()
            while "**" in line:
                before, bold, after = re.split(r"\*\*(.+?)\*\*", line, 1)
                paragraph.add_run(before)
                paragraph.add_run(bold).bold = True
                line = after
            paragraph.add_run(line)

    doc.save("output.docx")


def main():

    # Gemini API config
    genai.configure(api_key="AIzaSyAjLpSkv_e9a5rLf-Efl90l2JEkh7ZWzMI")

    text = extract_text()
    prompt = build_prompt(text=text)

    response = get_response(prompt)

    print("Starting to writing...\n")
    write_to_docx(response)
    # with open("result.txt", "w", encoding="utf-8") as file:
    #     file.write(response)
    #     print("Completed!")

    print(response)



if __name__ == "__main__":
    main()




